import json
import sys
import os
from dotenv import load_dotenv
from typing import Annotated, TypedDict, List, Dict, Sequence
from langchain_core.messages import BaseMessage
from langgraph.graph.message import add_messages
from langgraph.graph import StateGraph, START, END
from pydantic import  Field, create_model
from langchain_openai import ChatOpenAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from openai import OpenAI
from langchain_core.messages import HumanMessage, AIMessage
from docx import Document

load_dotenv()

os.environ["OPENAI_API_KEY"] = os.getenv("OPENAI_API_KEY")

topic = input("Enter the topic: ")

OPENAI_API_KEY = os.environ["OPENAI_API_KEY"]

class ReportGenerator(TypedDict):
  messages: Annotated[Sequence[BaseMessage], add_messages]
  outline: Dict[str, str]
  current_section: int
  section_content: str
  total_sections: int
  full_report: List[Dict[str, str]]

class Article(TypedDict):
  title: str
  type: str
  url: str

graph_builder = StateGraph(ReportGenerator)

llm = ChatOpenAI(model='gpt-4o-2024-08-06')

def create_outline_model(section_count: int):
  fields = {f"section{i}": (str, Field(description=f"Title for section {i}")) for i in range(1, section_count + 1)}
  return create_model("DynamicOutline", **fields)

fields = {f"section{i}": (str, Field(description=f"Title for section {i}")) for i in range(1, 3 + 1)}

DynamicOutline = create_model("DynamicOutline", **fields)

DynamicOutline = create_outline_model(20)

def outline_generator(state: ReportGenerator):
  DynamicOutline = create_outline_model(state["total_sections"])
  outline_parser = JsonOutputParser(pydantic_object=DynamicOutline)

  outline_prompt = PromptTemplate(
      template="""
      Create an outline for a detailed report with exactly {section_count} main sections.
      {format_instructions}
      The topic is: {topic}
      """,
      input_variables=["section_count", "topic"],
      partial_variables={"format_instructions": outline_parser.get_format_instructions()},
  )
  
  chain = outline_prompt | llm | outline_parser
  
  outline = chain.invoke({
    "section_count": state["total_sections"], 
    "topic": state["messages"][-1].content
  })
  return {"outline": outline}

client = OpenAI()

def content_writer(state: ReportGenerator):
  file_path = 'data.json'
  search_results = None
  
  with open(file_path, "r", encoding="utf-8") as file:
    search_results = json.load(file)

  if "error" in state:
    return {"messages": [AIMessage(content=f"An error occurred: {state['error']}")]}
  
  if state["current_section"] > state["total_sections"]:
    return {"messages": [AIMessage(content="Report completed.")]}
  
  current_section_key = f"section{state['current_section']}"
  current_topic = state["outline"][current_section_key]
  
  previous_sections_content = []
  for i in range(1, state['current_section']):
    section_key = f"section{i}"
    if section_key in state["section_content"]:
      previous_sections_content.append(f"""
      Section {i}: 
      {state['outline'][section_key]}
      {state['section_content'][section_key]}
      """)
  
  previous_sections = "\n\n".join(previous_sections_content)
  
  prompt = PromptTemplate(
    template="""
      Your role is write a professional and detailed report on the following topic: {topic}.
      You should understand well about the {topic}.

      Use the following search results to gather information: {search_results}

      It must be a detailed section with statistics and information.
      Also, It always related with Crypto Currency and Blockchain.
      Especially, Meme coin of Solana ecosystem.
      
      Previous sections:
      {previous_sections}
      Write only the content for this section, 
      do not include any image prompts or suggestions.
      Detailed statistics or information is needed, 
      so you should include collected information from search result

      This report will use for educational research for memecoin experted ai model.
    """,
    input_variables=["topic", "search_results", "previous_sections"]
  )
  section_content = llm.invoke(prompt.format(
    topic=topic,
    search_results=search_results,
    previous_sections=previous_sections
  ))

  return {
    "section_content": section_content.content,
    "current_section": state["current_section"] + 1,
    "full_report": state["full_report"] + [{
        "title": current_topic,
        "content": section_content.content
    }]
  }

def report_generator(state: ReportGenerator):
  file_path = 'data/report.docx'

  doc = Document()
  doc.add_heading(f"Report: {state['messages'][0].content}", 0)

  for section in state['full_report']:
    doc.add_heading(section['title'], level=1)
    doc.add_paragraph(section['content'])
    doc.add_page_break()

  doc.save(file_path)


def should_continue_writing(state: ReportGenerator):
  if state["current_section"] <= state["total_sections"]:
      return "write_section"
  else:
      return "finalize_report"

graph_builder.add_conditional_edges(
    "content_writer",
    should_continue_writing,
    {
      "write_section": "content_writer",
      "finalize_report": "report_generator"
    }
)


current_directory = os.getcwd()

def main(argv):
  if (len(argv) < 3):
    sys.exit(1)

  file_path = argv[1]
  total_sections = argv[2]
  try:
      graph_builder.add_node('outline_generator', outline_generator)
      graph_builder.add_node('content_writer', content_writer)
      graph_builder.add_node('report_generator', report_generator)

      graph_builder.add_edge(START, 'outline_generator')
      graph_builder.add_edge('outline_generator', 'content_writer')
      graph_builder.add_edge('report_generator', END)
      
      graph = graph_builder.compile()

      initial_state: ReportGenerator = {
        "messages": [HumanMessage(content=topic)],
        "total_sections": int(total_sections),
        "current_section": 1,
        "full_report": [],
      }

      for chunk in graph.stream(initial_state,stream_mode="update"):
        print(chunk)

  except FileNotFoundError:
    print(f"Error: {file_path} not found.")
    sys.exit(1)
  except json.JSONDecodeError as e:
    print(f"Error: {e}")
    sys.exit(1)

if __name__ == "__main__":
    main(sys.argv)